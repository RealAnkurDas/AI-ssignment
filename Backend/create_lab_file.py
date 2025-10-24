from utils.image_utils import create_output_image
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont
from dotenv import load_dotenv
from docx.shared import Inches
from docx.shared import Pt
from docx import Document
from io import StringIO
import tempfile
import openai
import sys
import os

import pandas

load_dotenv()

groq_api_key = os.getenv("GROQ_API_KEY")

class LabFile:
    @staticmethod
    def create_python_lab_file(student_name, aud_number, lab_number, questions):
        # Initialize groq connection
        client = openai.OpenAI(
                api_key=groq_api_key,
                base_url="https://api.groq.com/openai/v1"
            )
        
        # Make document object
        document = Document()
        
        # Main Title
        p = document.add_paragraph('')
        # Set alignment to center
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add run and set font size, bold, and underline
        run = p.add_run(f'Python Lab Week {str(lab_number)}')
        font = run.font
        font.size = Pt(16)
        font.bold = True
        font.underline = True

        # Solutions
        for qno, question in enumerate(questions):

            # Step 1. Generate code using groq
            code_guidelines = f"""
            Return only executable Python code, no other text.
            - No markdown code blocks or formatting
            - No explanatory text or comments
            - No headers or titles
            - No if __name__ == "__main__" blocks
            - Write code that runs directly when executed
            - For demonstration purposes, use hardcoded sample data instead of user input
            - You must include the Name as {student_name}
            - You must include the AUD as {aud_number}
            
            Example:
            print("Name: {student_name}")
            print("AUD: {aud_number}")
            def factorial(n):
                if n == 0:
                    return 1
                return n * factorial(n-1)
            
            result = factorial(5)
            print("Factorial:",result)
            """
            
            code_response = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[
                    {
                        "role": "system",
                        "content": f"You are a Python code generator. {code_guidelines}"
                    },
                    {
                        "role": "user",
                        "content": f"Write Python code for: {question}"
                    }
                ],
                temperature=0.7
            )

            code = code_response.choices[0].message.content.strip()
            
            # Remove markdown code blocks if present
            if code.startswith("```python"):
                code = code[9:]
            elif code.startswith("```"):
                code = code[3:]
            if code.endswith("```"):
                code = code[:-3]

            # Step 2. Generate appropriate output using groq
            output_guidelines = f"""
            Generate the expected output for the given Python code.
            - Return only the output text, no other formatting
            - For input-driven code, assume reasonable sample inputs
            - Show what the program would display when executed
            - Include any print statements, results, or program output
            - Make it realistic and appropriate for the code
            - Do not include code or explanations, just the output
            """
            
            output_response = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[
                    {
                        "role": "system",
                        "content": f"You are an output generator. {output_guidelines}"
                    },
                    {
                        "role": "user",
                        "content": f"Generate the expected output for this Python code:\n\n{code}"
                    }
                ],
                temperature=0.7
            )

            output = output_response.choices[0].message.content.strip()

            # Subtitle: Question
            document.add_paragraph('', style='List Number').add_run(str(question)).bold = True

            # Subtitle: Code
            document.add_paragraph('').add_run('Code:').bold = True

            # Code
            document.add_paragraph(code)

            # Subtitle: Output
            document.add_paragraph('').add_run('Output:').bold = True

            # Output Image
            temp_image_path = create_output_image(output)
            try:
                # Add the image to document with appropriate width
                document.add_picture(temp_image_path, width=Inches(6))
            finally:
                # Clean up temporary image file
                if os.path.exists(temp_image_path):
                    os.unlink(temp_image_path)
            
            # New Page
            document.add_page_break()

            print(f"Processed question {qno+1}...")

        # Create output directory if it doesn't exist
        output_dir = 'output'
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        filename = f'{student_name} Python Lab Week {lab_number}.docx'
        filepath = os.path.join(output_dir, filename)
        
        # Save the document
        document.save(filepath)
        
        return filepath

if __name__ == "__main__":
    q = ["WAP to print the first 10 natural numbers", 
    "WAP to find factorial of n", 
    "WAP to do recursive quicksort"]

    LabFile().create_python_lab_file("John Doe", 12345, 1, q)
